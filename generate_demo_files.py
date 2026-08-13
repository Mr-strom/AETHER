"""
generate_demo_files.py
Run this in your conda environment to create realistic test documents
for the AETHER ingestion pipeline.

Usage:
    conda activate aether
    my_path = r"C:/Users/pkuma/projects/AETHER"
    python generate_demo_files.py
"""

import os
from pathlib import Path

# Create demo bundle directory
DEMO_DIR = Path("./demo_bundle")
DEMO_DIR.mkdir(exist_ok=True)

# ============================================================
# 1. TXT FILE — Field Notes
# ============================================================
txt_content = """AETHER Project — Field Inspection Notes
==========================================

Date: 2024-03-15
Inspector: J. Doe
Location: Building 7, Floor 3

## Pre-Inspection Checklist

- [x] Safety helmet worn
- [x] Voltage tester calibrated
- [ ] Inspection form printed

## Panel A-001 Inspection

The terminal block on Panel A-001 shows signs of corrosion on screws 2 and 4.
Voltage reading: 112V (within acceptable range 110-120V).
No overheating detected. Thermal imaging normal.

Recommended action: Replace terminal screws 2 and 4 during next maintenance window.

## Panel B-002 Inspection

Panel B-002 passed all checks. Voltage: 118V. Clean terminals.
No anomalies detected.

## Voice Note Summary

Voice note from 2024-03-15 14:30:
"The corrosion on Panel A is worse than last month. We need to schedule replacement
before the rainy season starts. The voltage fluctuated between 111 and 113V during
testing. Document everything for the compliance report."

## Missing Information

- Serial number of Panel A-001 not visible (paint cover)
- Last maintenance date unknown
- Manufacturer documentation not on-site
"""

with open(DEMO_DIR / "field_notes.txt", "w", encoding="utf-8") as f:
    f.write(txt_content)

print("✅ Created: demo_bundle/field_notes.txt")

# ============================================================
# 2. CSV FILE — Equipment Inventory
# ============================================================
csv_content = """equipment_id,location,install_date,last_inspection,status,voltage_reading,inspector
PANEL-A-001,Building-7-Floor-3,2022-01-10,2024-03-15,CONDITIONAL,112,J.Doe
PANEL-B-002,Building-7-Floor-3,2022-01-10,2024-03-15,PASS,118,J.Doe
PANEL-C-003,Building-7-Floor-2,2021-08-15,2024-02-20,PASS,115,M.Smith
PANEL-D-004,Building-7-Floor-2,2021-08-15,2024-02-20,FAIL,98,M.Smith
TRANSFORMER-T1,Building-7-Basement,2020-03-20,2024-01-10,PASS,480,A.Johnson
UPS-UNIT-01,Building-7-Server-Room,2023-06-01,2024-03-10,PASS,230,R.Williams
"""

with open(DEMO_DIR / "equipment_inventory.csv", "w", encoding="utf-8") as f:
    f.write(csv_content)

print("✅ Created: demo_bundle/equipment_inventory.csv")

# ============================================================
# 3. MD FILE — Maintenance Manual (Markdown)
# ============================================================
md_content = """# Electrical Panel Maintenance Manual v2.4

## Table of Contents

1. Safety Procedures
2. Inspection Steps
3. Voltage Standards
4. Troubleshooting Guide
5. Compliance Requirements

---

## 1. Safety Procedures

### 1.1 Personal Protective Equipment

All personnel must wear:
- Insulated gloves (Class 00 minimum)
- Safety helmet with face shield
- Arc-rated clothing (minimum 8 cal/cm²)

### 1.2 Lockout/Tagout

Before opening any panel:
1. De-energize the circuit
2. Apply lockout device
3. Verify zero energy state with calibrated tester
4. Tag the panel with date and personnel ID

---

## 2. Inspection Steps

### Step 1 — Visual Inspection

Check for:
- Physical damage to enclosure
- Signs of overheating (discoloration, melting)
- Corrosion on terminals and bus bars
- Proper grounding connections

### Step 2 — Voltage Measurement

Use calibrated digital multimeter:
- Phase-to-phase voltage: 380V ± 5%
- Phase-to-neutral voltage: 220V ± 5%
- Record all readings in inspection form

### Step 3 — Terminal Tightness

Torque specifications:
- M6 screws: 5.5 Nm
- M8 screws: 13.5 Nm
- M10 screws: 27.0 Nm

### Step 4 — Thermal Imaging

Scan all connections with IR camera:
- Normal: ΔT < 10°C above ambient
- Caution: ΔT 10-20°C — schedule follow-up
- Critical: ΔT > 20°C — immediate action required

---

## 3. Voltage Standards

| Parameter | Acceptable Range | Critical Threshold |
|-----------|-----------------|-------------------|
| Phase-Phase | 361-399 V | < 340 or > 420 V |
| Phase-Neutral | 209-231 V | < 198 or > 242 V |
| Frequency | 49.5-50.5 Hz | < 48 or > 52 Hz |

---

## 4. Troubleshooting Guide

### Symptom: Voltage Drop Under Load

**Possible Causes:**
1. Loose terminal connections
2. Undersized conductors
3. High impedance in distribution path

**Diagnostic Steps:**
1. Measure voltage at panel input (no load)
2. Measure voltage at panel input (full load)
3. Calculate voltage drop percentage
4. If drop > 3%, inspect all connections upstream

### Symptom: Terminal Corrosion

**Possible Causes:**
1. Humidity ingress
2. Chemical exposure
3. Galvanic reaction between dissimilar metals

**Remediation:**
1. De-energize and lock out
2. Remove corroded hardware
3. Clean with electrical contact cleaner
4. Apply antioxidant compound
5. Replace with compatible materials

---

## 5. Compliance Requirements

### NFPA 70E (Electrical Safety)

All inspections must comply with NFPA 70E Article 130:
- Arc flash hazard analysis updated every 5 years
- PPE category determined by incident energy calculation
- Approach boundaries clearly marked

### OSHA 1910.147 (Lockout/Tagout)

- Written procedure required for each panel
- Annual training for all authorized personnel
- Audit of procedures every 12 months

### Internal Standards

- Inspection frequency: quarterly for critical panels
- Documentation retention: 7 years minimum
- Digital photos required for all findings
- Inspector certification: Level 2 minimum

---

*Document Control: Revision 2.4 — Approved by Chief Engineer on 2024-01-15*
"""

with open(DEMO_DIR / "maintenance_manual.md", "w", encoding="utf-8") as f:
    f.write(md_content)

print("✅ Created: demo_bundle/maintenance_manual.md")

# ============================================================
# 4. PDF FILE — Inspection Report (using fpdf2)
# ============================================================
# ============================================================
# 4. PDF FILE — Inspection Report
# ============================================================
try:
    from fpdf import FPDF
    
    class PDFReport(FPDF):
        def header(self):
            self.set_font("helvetica", "B", 16)
            self.cell(0, 10, "Quarterly Inspection Report - Q1 2024", new_x="LMARGIN", new_y="NEXT", align="C")
            self.ln(5)
        
        def chapter_title(self, title):
            self.set_font("helvetica", "B", 14)
            self.set_fill_color(230, 230, 230)
            self.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT", fill=True)
            self.ln(2)
        
        def chapter_body(self, body):
            self.set_font("helvetica", "", 11)
            self.multi_cell(0, 6, body)
            self.ln()
    
    pdf = PDFReport()
    pdf.add_page()
    
    pdf.chapter_title("Executive Summary")
    pdf.chapter_body(
        "This report covers the quarterly inspection of electrical panels in Building 7. "
        "A total of 6 panels were inspected. One panel (PANEL-A-001) requires conditional "
        "maintenance due to terminal corrosion. All other panels passed inspection."
    )
    
    pdf.chapter_title("Panel A-001 - Detailed Findings")
    pdf.chapter_body(
        "Location: Building 7, Floor 3, East Wing\n"
        "Inspector: J. Doe\n"
        "Date: March 15, 2024\n\n"
        "FINDINGS:\n"
        "1. Terminal screws 2 and 4 show visible corrosion (green oxidation).\n"
        "2. Voltage reading: 112V (acceptable per standard 110-120V).\n"
        "3. No thermal anomalies detected (max delta-T: 4C).\n"
        "4. Grounding impedance: 0.8 ohms (within < 1.0 ohm requirement).\n\n"
        "RECOMMENDATION:\n"
        "Schedule terminal screw replacement during next planned outage. "
        "Estimated downtime: 2 hours. Replacement parts: M6 stainless steel screws, "
        "antioxidant compound, contact cleaner."
    )
    
    pdf.chapter_title("Panel B-002 - Detailed Findings")
    pdf.chapter_body(
        "Location: Building 7, Floor 3, West Wing\n"
        "Inspector: J. Doe\n"
        "Date: March 15, 2024\n\n"
        "FINDINGS:\n"
        "1. All terminals clean and properly torqued.\n"
        "2. Voltage reading: 118V (nominal).\n"
        "3. Thermal imaging: normal across all connections.\n"
        "4. No physical damage to enclosure.\n\n"
        "RECOMMENDATION:\n"
        "No action required. Next inspection scheduled for June 2024."
    )
    
    pdf.chapter_title("Compliance Checklist")
    pdf.chapter_body(
        "[X] NFPA 70E Article 130 compliance verified\n"
        "[X] OSHA 1910.147 LOTO procedures followed\n"
        "[X] Inspector certification valid (Level 2, expires 2025-08)\n"
        "[X] Calibration certificates current for all test equipment\n"
        "[ ] Serial number for Panel A-001 not documented (paint obstruction)\n"
        "[ ] Previous maintenance records missing for Panel C-003"
    )
    
    pdf.output(str(DEMO_DIR / "inspection_report_q1_2024.pdf"))
    print("Created: demo_bundle/inspection_report_q1_2024.pdf")
    
except Exception as e:
    print(f"PDF generation failed: {e}")
# ============================================================
# 5. DOCX FILE — Incident Report (using python-docx)
# ============================================================
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Incident Report — Voltage Fluctuation Event", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    cells = table.rows[0].cells
    cells[0].text = "Incident ID"
    cells[1].text = "INC-2024-0315-001"
    cells = table.rows[1].cells
    cells[0].text = "Date/Time"
    cells[1].text = "2024-03-15 09:45 AM"
    cells = table.rows[2].cells
    cells[0].text = "Location"
    cells[1].text = "Building 7, Floor 3, Panel A-001"
    cells = table.rows[3].cells
    cells[0].text = "Reporter"
    cells[1].text = "J. Doe (Level 2 Inspector)"

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "On March 15, 2024, at approximately 09:45 AM, voltage fluctuations were detected "
        "on Panel A-001 in Building 7. The voltage dropped from nominal 115V to 108V "
        "over a 3-minute period before stabilizing. No equipment damage was reported. "
        "Root cause analysis points to corroded terminal screws creating intermittent "
        "high-resistance connections."
    )

    # Timeline
    doc.add_heading("2. Event Timeline", level=1)
    timeline = [
        ("09:42", "Normal voltage reading: 115V"),
        ("09:45", "Voltage drop detected: 108V"),
        ("09:46", "Alert triggered in monitoring system"),
        ("09:48", "Inspector dispatched to location"),
        ("10:05", "Visual inspection completed"),
        ("10:15", "Terminal corrosion identified as likely cause"),
        ("10:30", "Temporary monitoring equipment installed"),
        ("11:00", "Voltage stabilized at 112V"),
    ]
    for time, event in timeline:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{time} — ").bold = True
        p.add_run(event)

    # Findings
    doc.add_heading("3. Technical Findings", level=1)
    doc.add_heading("3.1 Voltage Analysis", level=2)
    doc.add_paragraph(
        "The voltage drop from 115V to 108V represents a 6.1% deviation from nominal. "
        "Per maintenance manual Section 3, the acceptable range is 209-231V (phase-neutral) "
        "or 110-120V for this specific panel type. While 108V is below the acceptable threshold, "
        "the duration was insufficient to trigger protective relay operation."
    )

    doc.add_heading("3.2 Physical Inspection", level=2)
    doc.add_paragraph(
        "Terminal screws 2 and 4 on the main breaker showed significant green oxidation. "
        "Torque measurement on screw 2: 3.2 Nm (specification: 5.5 Nm). Screw 4 was "
        "loose enough to rotate by hand. These conditions explain the intermittent "
        "high-resistance connection causing voltage sag under load."
    )

    # Contradictory evidence (for testing conflict detection later)
    doc.add_heading("3.3 Contradictory Assessment", level=2)
    doc.add_paragraph(
        "NOTE: The automated monitoring system log (from SCADA) shows the voltage fluctuation "
        "started at 09:41, not 09:45 as recorded by the on-site inspector. The SCADA timestamp "
        "is considered authoritative. Inspector report timestamp should be corrected."
    )

    # Recommendations
    doc.add_heading("4. Recommendations", level=1)
    recs = [
        "Immediate: Replace terminal screws 2 and 4 on Panel A-001",
        "Immediate: Verify torque on all remaining terminals (M6: 5.5 Nm)",
        "Short-term: Install continuous voltage monitoring on Panel A-001",
        "Short-term: Review SCADA timestamp synchronization across all Building 7 panels",
        "Long-term: Replace all carbon steel screws with stainless steel in humid environments",
        "Long-term: Reduce inspection interval for Panel A-001 from quarterly to monthly",
    ]
    for rec in recs:
        doc.add_paragraph(rec, style="List Number")

    # Sign-off
    doc.add_paragraph()
    doc.add_paragraph("_" * 40)
    p = doc.add_paragraph("Inspector Signature: J. Doe")
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph("Date: 2024-03-15")
    p.runs[0].font.size = Pt(10)

    doc.save(str(DEMO_DIR / "incident_report.docx"))
    print("✅ Created: demo_bundle/incident_report.docx")

except ImportError:
    print("⚠️  python-docx not installed. Installing...")
    os.system("pip install python-docx")
    print("   Please re-run this script to generate the DOCX.")

# ============================================================
# Summary
# ============================================================
print()
print("=" * 60)
print("DEMO BUNDLE CREATED SUCCESSFULLY")
print("=" * 60)
print(f"Location: {DEMO_DIR.absolute()}")
print()
print("Files generated:")
for f in sorted(DEMO_DIR.iterdir()):
    size = f.stat().st_size
    print(f"  📄 {f.name:<40} {size:>10,} bytes")
print()
print("Next steps:")
print("  1. Test ingestion: POST /api/sources/import with path=./demo_bundle")
print("  2. Check that all 4 file types are parsed correctly")
print("  3. Verify chunks are created with correct metadata")
print("  4. Test embedding + FAISS indexing on the chunks")
print()