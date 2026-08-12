"""python-docx based DOCX ingestion extractor.

Extracts paragraphs (preserving Heading levels as Markdown ``#`` prefixes)
and tables (rendered as a Markdown table) from ``.docx`` files.

Heading mapping
---------------
``Heading 1`` → ``# …``
``Heading 2`` → ``## …``
``Heading 3``–``Heading 9`` → ``### …``
Any other paragraph style → plain text.

Table rendering
---------------
Each table is converted to a GitHub-flavoured Markdown table with the
first row used as headers.  The table is then chunked the same way as
regular text if it exceeds ``chunk_size``.

Example
-------
>>> import asyncio
>>> chunks = asyncio.run(DocxIngester().parse("document.docx"))
>>> print(chunks[0].text[:60])
# Executive Summary  This section summarises the findings…
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import List

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
except ImportError:
    DocxDocument = None  # type: ignore
    qn = None  # type: ignore

from backend.app.config import settings
from backend.services.ingest.schema import IngestChunk
from backend.services.ingest.text import _split_into_chunks

logger = logging.getLogger(__name__)

_HEADING_PREFIX: dict[str, str] = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "### ",
    "Heading 5": "### ",
    "Heading 6": "### ",
    "Heading 7": "### ",
    "Heading 8": "### ",
    "Heading 9": "### ",
}


def _para_to_text(para: "docx.text.paragraph.Paragraph") -> str:  # type: ignore[name-defined]
    """Convert a docx paragraph to Markdown-prefixed plain text.

    Args:
        para: A ``python-docx`` Paragraph object.

    Returns:
        Marked-up text string; empty string if the paragraph is blank.
    """
    raw = para.text.strip()
    if not raw:
        return ""
    prefix = _HEADING_PREFIX.get(para.style.name, "")
    return f"{prefix}{raw}"


def _table_to_markdown(table: "docx.table.Table") -> str:  # type: ignore[name-defined]
    """Render a docx table as a Markdown table string.

    The first row is treated as the header row.

    Args:
        table: A ``python-docx`` Table object.

    Returns:
        Multi-line string in GitHub-flavoured Markdown table format.
        Returns an empty string if the table has no rows.
    """
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return ""

    col_count = max(len(r) for r in rows)
    # Pad short rows
    rows = [r + [""] * (col_count - len(r)) for r in rows]

    header = rows[0]
    sep = ["---"] * col_count
    body = rows[1:]

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


class DocxIngester:
    """python-docx backed DOCX extractor.

    Parameters
    ----------
    chunk_size:
        Maximum characters per chunk.
    overlap:
        Overlap characters between consecutive chunks.
    """

    def __init__(
        self,
        chunk_size: int | None = None,
        overlap: int | None = None,
    ) -> None:
        self.chunk_size = chunk_size or settings.CHUNK_SIZE_CHARS
        self.overlap = overlap or settings.CHUNK_OVERLAP_CHARS

    async def parse(self, file_path: str | Path) -> List[IngestChunk]:
        """Extract text from a ``.docx`` file into evidence chunks.

        Paragraphs and tables are interleaved in document order.  Each
        element is accumulated into a running buffer; when the buffer
        reaches ``chunk_size``, it is flushed into a new chunk.

        Args:
            file_path: Path to the ``.docx`` file.

        Returns:
            Ordered list of :class:`IngestChunk` objects.

        Raises:
            FileNotFoundError: If *file_path* does not exist.
            RuntimeError: If python-docx cannot open the document.
        """
        if DocxDocument is None:
            raise RuntimeError("python-docx is required: pip install python-docx")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        logger.info("Ingesting DOCX: %s", path)

        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            logger.error("python-docx failed to open '%s': %s", path, exc)
            raise RuntimeError(f"Cannot open DOCX '{path}': {exc}") from exc

        # Collect doc body elements in order (paragraphs AND tables).
        # python-docx iterates doc.paragraphs / doc.tables separately,
        # losing interleaving.  We walk the XML body directly instead.
        body_elements: List[str] = []

        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                # Find the matching paragraph object
                for para in doc.paragraphs:
                    if para._element is child:
                        text = _para_to_text(para)
                        if text:
                            body_elements.append(text)
                        break
            elif tag == "tbl":
                for tbl in doc.tables:
                    if tbl._tbl is child:
                        md = _table_to_markdown(tbl)
                        if md:
                            body_elements.append(md)
                        break

        full_text = "\n\n".join(body_elements)

        chunks = _split_into_chunks(
            text=full_text,
            chunk_size=self.chunk_size,
            overlap=self.overlap,
            source_path=str(path),
        )

        logger.info("Extracted %d chunks from DOCX '%s'", len(chunks), path.name)
        return chunks


docx_ingester = DocxIngester()
