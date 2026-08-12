"""Unit tests for backend/services/ingest/docx.py"""

import asyncio
from pathlib import Path
import pytest

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from backend.services.ingest.schema import IngestChunk

pytestmark = pytest.mark.skipif(
    not DOCX_AVAILABLE, reason="python-docx not installed"
)


def _make_docx(path: Path, paragraphs: list[tuple[str, str]]) -> None:
    doc = Document()
    for style, text in paragraphs:
        if style == "Normal":
            doc.add_paragraph(text)
        else:
            doc.add_heading(text, level=int(style.split()[-1]) if style[-1].isdigit() else 1)
    doc.save(str(path))


def test_docx_basic(tmp_path: Path) -> None:
    from backend.services.ingest.docx import DocxIngester

    path = tmp_path / "basic.docx"
    _make_docx(path, [("Normal", "Evidence paragraph one."), ("Normal", "Evidence paragraph two.")])

    ingester = DocxIngester()
    chunks = asyncio.run(ingester.parse(path))

    assert len(chunks) >= 1
    combined = " ".join(c.text for c in chunks)
    assert "Evidence paragraph one" in combined
    assert "Evidence paragraph two" in combined


def test_docx_heading_prefixes(tmp_path: Path) -> None:
    from backend.services.ingest.docx import DocxIngester

    path = tmp_path / "headings.docx"
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Body paragraph under introduction.")
    doc.add_heading("Section Two", level=2)
    doc.save(str(path))

    ingester = DocxIngester()
    chunks = asyncio.run(ingester.parse(path))

    combined = " ".join(c.text for c in chunks)
    assert "Introduction" in combined
    assert "Section Two" in combined


def test_docx_table_rendered(tmp_path: Path) -> None:
    from backend.services.ingest.docx import DocxIngester

    path = tmp_path / "table.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "30"
    table.cell(2, 0).text = "Bob"
    table.cell(2, 1).text = "25"
    doc.save(str(path))

    ingester = DocxIngester()
    chunks = asyncio.run(ingester.parse(path))

    combined = " ".join(c.text for c in chunks)
    assert "Name" in combined
    assert "Alice" in combined


def test_docx_empty_document(tmp_path: Path) -> None:
    from backend.services.ingest.docx import DocxIngester

    path = tmp_path / "empty.docx"
    Document().save(str(path))

    ingester = DocxIngester()
    chunks = asyncio.run(ingester.parse(path))
    assert chunks == []


def test_docx_not_found() -> None:
    from backend.services.ingest.docx import DocxIngester

    ingester = DocxIngester()
    with pytest.raises(FileNotFoundError):
        asyncio.run(ingester.parse("/no/such/file.docx"))


def test_docx_chunk_indices_sequential(tmp_path: Path) -> None:
    from backend.services.ingest.docx import DocxIngester

    path = tmp_path / "long.docx"
    doc = Document()
    for i in range(100):
        doc.add_paragraph(f"This is evidence sentence number {i}. " * 10)
    doc.save(str(path))

    ingester = DocxIngester(chunk_size=512, overlap=50)
    chunks = asyncio.run(ingester.parse(path))

    assert len(chunks) > 1
    assert [c.chunk_index for c in chunks] == list(range(len(chunks)))


def test_docx_source_path_set(tmp_path: Path) -> None:
    from backend.services.ingest.docx import DocxIngester

    path = tmp_path / "sp.docx"
    _make_docx(path, [("Normal", "Source path test.")])

    ingester = DocxIngester()
    chunks = asyncio.run(ingester.parse(path))

    for c in chunks:
        assert c.source_path == str(path)
